import streamlit as st
import subprocess
import os
import tempfile
import shutil
import uuid
from pathlib import Path
import platform
import time
import random
import docx
from docx.oxml.ns import qn
from xml.sax.saxutils import escape
import openpyxl

# ---------------------------------------------------------------------------
# SETUP DO PATH DO LIBREOFFICE
# ---------------------------------------------------------------------------
def setup_libreoffice_path():
    """Adiciona o LibreOffice ao PATH do sistema se necessário."""
    if platform.system() == "Windows":
        possible_paths = [
            r"C:\Program Files\LibreOffice\program",
            r"C:\Program Files (x86)\LibreOffice\program",
        ]
    elif platform.system() == "Darwin":
        possible_paths = ["/Applications/LibreOffice.app/Contents/MacOS"]
    else:
        possible_paths = ["/usr/bin", "/usr/local/bin"]

    for path in possible_paths:
        if os.path.exists(path) and path not in os.environ.get("PATH", ""):
            os.environ["PATH"] = path + os.pathsep + os.environ.get("PATH", "")
            break

setup_libreoffice_path()

# ---------------------------------------------------------------------------
# CONFIGURAÇÃO DE PÁGINA E CSS (EXATAMENTE O SEU CSS COMPACTO DE ZERO ESPAÇO)
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="Conversor de Documentos",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
    .main {
        background-color: #ffffff;
        color: #333333;
    }
    .block-container {
        padding-top: 0rem !important;
        padding-bottom: 0rem !important;
        max-width: 28rem !important;
    }
    header {display: none !important;}
    footer {display: none !important;}
    #MainMenu {display: none !important;}
    div[data-testid="stAppViewBlockContainer"] {
        padding-top: 0 !important;
        padding-bottom: 0 !important;
    }
    div[data-testid="stVerticalBlock"] {
        gap: 0 !important;
        padding-top: 0 !important;
        padding-bottom: 0 !important;
    }
    .element-container {
        margin-top: 0 !important;
        margin-bottom: 0 !important;
    }
    .stDownloadButton button {
        width: 100% !important;
        padding: 0.6rem 2rem;
    }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# TABELA DE NORMALIZAÇÃO DE NOMES POSTSCRIPT DE FONTES
# ---------------------------------------------------------------------------
POSTSCRIPT_TO_FAMILY = {
    'ArialMT': 'Arial',
    'Arial-BoldMT': 'Arial',
    'Arial-ItalicMT': 'Arial',
    'Arial-BoldItalicMT': 'Arial',
    'ArialBlack': 'Arial Black',
    'Verdana-Bold': 'Verdana',
    'Verdana-Italic': 'Verdana',
    'Verdana-BoldItalic': 'Verdana',
    'TimesNewRomanPSMT': 'Times New Roman',
    'TimesNewRomanPS-BoldMT': 'Times New Roman',
    'TimesNewRomanPS-ItalicMT': 'Times New Roman',
    'TimesNewRomanPS-BoldItalicMT': 'Times New Roman',
    'CourierNewPSMT': 'Courier New',
    'CourierNewPS-BoldMT': 'Courier New',
    'CourierNewPS-ItalicMT': 'Courier New',
    'CourierNewPS-BoldItalicMT': 'Courier New',
    'Calibri-Bold': 'Calibri',
    'Calibri-Italic': 'Calibri',
    'Calibri-Light': 'Calibri Light',
    'Tahoma-Bold': 'Tahoma',
}

# ---------------------------------------------------------------------------
# SANITIZAÇÃO DE XML DO DOCX
# ---------------------------------------------------------------------------
def sanitize_docx_xml(input_docx_path: str, output_docx_path: str):
    """
    Higieniza a estrutura interna do DOCX antes de passar para o LibreOffice.
    """
    try:
        doc = docx.Document(input_docx_path)

        def clean_rFonts(rPr):
            if rPr is None:
                return
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                for attr in ['asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'csTheme', 'cstheme']:
                    if qn(f'w:{attr}') in rFonts.attrib:
                        del rFonts.attrib[qn(f'w:{attr}')]

                for font_key in ['ascii', 'hAnsi', 'cs']:
                    val = rFonts.attrib.get(qn(f'w:{font_key}'))
                    if val:
                        if val in POSTSCRIPT_TO_FAMILY:
                            val = POSTSCRIPT_TO_FAMILY[val]
                        elif '-' in val:
                            base_name = val.split('-')[0]
                            if base_name in ['Verdana', 'Arial', 'Calibri', 'Tahoma', 'Georgia', 'Cambria', 'Helvetica']:
                                val = base_name
                        rFonts.attrib[qn(f'w:{font_key}')] = val

        for style in doc.styles:
            if hasattr(style, '_element') and style._element is not None:
                clean_rFonts(style._element.find(qn('w:rPr')))

        def process_paragraph(p):
            pPr = p._p.get_or_add_pPr()
            clean_rFonts(pPr.find(qn('w:rPr')))
            for run in p.runs:
                rPr_run = run._r.get_or_add_rPr()
                clean_rFonts(rPr_run)

        for p in doc.paragraphs:
            process_paragraph(p)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        process_paragraph(p)

        doc.save(output_docx_path)
        return True
    except Exception:
        shutil.copy2(input_docx_path, output_docx_path)
        return False

# ---------------------------------------------------------------------------
# HIGIENIZAÇÃO E LAYOUT AUTOMÁTICO DE PLANILHAS EXCEL (XLSX)
# ---------------------------------------------------------------------------
def sanitize_xlsx_page_setup(input_xlsx_path: str, output_xlsx_path: str):
    """
    Ajusta a orientação da planilha para Paisagem (Horizontal), ajusta a largura 
    para 1 página, corrige a altura das linhas para o tamanho da fonte (evitando corte do título)
    e define margens de segurança.
    """
    try:
        wb = openpyxl.load_workbook(input_xlsx_path)
        for ws in wb.worksheets:
            # 1. Configurações de Orientação e Papel (Paisagem / A4)
            ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
            ws.page_setup.paperSize = ws.PAPERSIZE_A4
            
            # 2. Ativa o ajuste de página para caber na largura de 1 página
            ws.sheet_properties.pageSetUpPr.fitToPage = True
            ws.page_setup.fitToWidth = 1
            ws.page_setup.fitToHeight = 0  # 0 permite que as linhas fluam livremente em várias páginas
            
            # 3. Define margens superiores e de cabeçalho confortáveis
            ws.page_margins.top = 0.75
            ws.page_margins.bottom = 0.5
            ws.page_margins.left = 0.4
            ws.page_margins.right = 0.4
            ws.page_margins.header = 0.3
            ws.page_margins.footer = 0.3

            # 4. Ajusta dinamicamente a altura de cada linha para garantir que fontes grandes (como 16pt no título)
            # não fiquem presas em alturas customizadas pequenas (como 14.25pt)
            for r in range(1, ws.max_row + 1):
                max_font_size = 0
                for c in range(1, ws.max_column + 1):
                    cell = ws.cell(row=r, column=c)
                    if cell.value is not None and cell.font and cell.font.size:
                        if cell.font.size > max_font_size:
                            max_font_size = cell.font.size
                
                if max_font_size > 0:
                    needed_height = max_font_size * 1.35 + 5  # Dá uma folga de segurança para o topo das letras
                    current_height = ws.row_dimensions[r].height or 15
                    if current_height < needed_height:
                        ws.row_dimensions[r].height = round(needed_height, 2)

        wb.save(output_xlsx_path)
        return True
    except Exception:
        shutil.copy2(input_xlsx_path, output_xlsx_path)
        return False

# ---------------------------------------------------------------------------
# GERENCIAMENTO DE FONTES E FONTCONFIG COM TTL CACHE (24 HORAS)
# ---------------------------------------------------------------------------
STATIC_FONTS_DIR = Path(__file__).resolve().parent / "static" / "fonts"
FONT_EXTENSIONS = {".ttf", ".otf", ".ttc", ".otc"}

@st.cache_resource(ttl=86400, show_spinner=False)
def prepare_font_environment():
    """
    Executa a cópia de fontes e fontconfig 1x a cada 24h.
    """
    user_fonts = Path.home() / ".fonts"
    user_share_fonts = Path.home() / ".local" / "share" / "fonts"
    
    for d in [user_fonts, user_share_fonts]:
        d.mkdir(parents=True, exist_ok=True)

    if STATIC_FONTS_DIR.is_dir():
        for font_path in STATIC_FONTS_DIR.rglob("*"):
            if font_path.is_file() and font_path.suffix.lower() in FONT_EXTENSIONS:
                for target_dir in [user_fonts, user_share_fonts]:
                    dest = target_dir / font_path.name
                    if not dest.exists() or dest.stat().st_size != font_path.stat().st_size:
                        try:
                            shutil.copy2(font_path, dest)
                        except OSError:
                            pass

    if platform.system() == "Linux":
        try:
            subprocess.run(["fc-cache", "-fv", str(STATIC_FONTS_DIR), str(user_fonts)], capture_output=True, timeout=15)
        except Exception:
            pass

    config_dir = Path.home() / ".config" / "fontconfig"
    config_dir.mkdir(parents=True, exist_ok=True)
    conf_file = config_dir / "fonts.conf"

    xml_lines = [
        '<?xml version="1.0"?>',
        '<!DOCTYPE fontconfig SYSTEM "fonts.dtd">',
        '<fontconfig>',
        '  <include ignore_missing="yes">/etc/fonts/fonts.conf</include>',
        f'  <dir>{escape(str(STATIC_FONTS_DIR))}</dir>',
        f'  <dir>{escape(str(user_fonts))}</dir>',
        f'  <dir>{escape(str(user_share_fonts))}</dir>',
        '',
        '  <!-- Aliases para direcionar fontes PostScript para as fontes reais -->',
    ]

    sans_family_maps = [
        ("ArialMT", ["Arial", "Liberation Sans", "DejaVu Sans", "sans-serif"]),
        ("Arial", ["Arial", "Liberation Sans", "DejaVu Sans", "sans-serif"]),
        ("Verdana-Bold", ["Verdana", "DejaVu Sans", "sans-serif"]),
        ("Verdana", ["Verdana", "DejaVu Sans", "Liberation Sans", "sans-serif"]),
        ("Helvetica", ["Arial", "Liberation Sans", "sans-serif"]),
        ("Calibri", ["Calibri", "Carlito", "Liberation Sans", "sans-serif"]),
        ("Tahoma", ["Tahoma", "DejaVu Sans", "sans-serif"]),
        ("Segoe UI", ["Segoe UI", "DejaVu Sans", "sans-serif"]),
    ]

    for source_font, target_list in sans_family_maps:
        xml_lines.append('  <alias>')
        xml_lines.append(f'    <family>{escape(source_font)}</family>')
        xml_lines.append('    <prefer>')
        for tgt in target_list:
            xml_lines.append(f'      <family>{escape(tgt)}</family>')
        xml_lines.append('    </prefer>')
        xml_lines.append('  </alias>')

    xml_lines.extend([
        '  <!-- Redirecionamento global: impede que Sans-serif caia em Serif -->',
        '  <match target="pattern">',
        '    <test qual="all" name="family" compare="not_eq"><string>serif</string></test>',
        '    <test qual="all" name="family" compare="not_eq"><string>monospace</string></test>',
        '    <test qual="all" name="family" compare="not_eq"><string>sans-serif</string></test>',
        '    <edit name="family" mode="append" binding="weak">',
        '      <string>sans-serif</string>',
        '    </edit>',
        '  </match>',
        '</fontconfig>'
    ])

    conf_file.write_text("\n".join(xml_lines), encoding="utf-8")

    env = os.environ.copy()
    env["FONTCONFIG_FILE"] = str(conf_file)
    env["FONTCONFIG_PATH"] = str(config_dir)
    env["SAL_VCL_FONTPATH"] = f"{STATIC_FONTS_DIR}:{user_fonts}:{user_share_fonts}"
    return env

# ---------------------------------------------------------------------------
# REGISTRO DE SUBSTITUIÇÃO INTERNO DO LIBREOFFICE
# ---------------------------------------------------------------------------
def _build_lo_registry_substitutions():
    rules = [
        ("ArialMT", "Arial"),
        ("Verdana-Bold", "Verdana"),
        ("Verdana", "Verdana"),
        ("Helvetica", "Arial"),
        ("Calibri", "Calibri"),
        ("Tahoma", "DejaVu Sans"),
        ("TimesNewRomanPSMT", "Times New Roman"),
        ("CourierNewPSMT", "Courier New"),
    ]
    items = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<oor:items xmlns:oor="http://openoffice.org/2001/registry" xmlns:xs="http://www.w3.org/2001/XMLSchema" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">',
        '<item oor:path="/org.openoffice.Office.Common/Font/Substitution"><prop oor:name="Replacement" oor:op="fuse"><value>true</value></prop></item>',
    ]
    for i, (replace_f, sub_f) in enumerate(rules):
        items.append(
            f'<item oor:path="/org.openoffice.Office.Common/Font/Substitution/FontPairs"><node oor:name="_{i}" oor:op="replace">'
            f'<prop oor:name="Always" oor:op="fuse"><value>true</value></prop>'
            f'<prop oor:name="OnScreenOnly" oor:op="fuse"><value>false</value></prop>'
            f'<prop oor:name="ReplaceFont" oor:op="fuse"><value>{escape(replace_f)}</value></prop>'
            f'<prop oor:name="SubstituteFont" oor:op="fuse"><value>{escape(sub_f)}</value></prop>'
            '</node></item>'
        )
    items.append("</oor:items>")
    return "\n".join(items)

def inject_lo_profile_settings(profile_dir: Path):
    try:
        profile_dir.mkdir(parents=True, exist_ok=True)
        (profile_dir / "registrymodifications.xcu").write_text(_build_lo_registry_substitutions(), encoding="utf-8")
    except OSError:
        pass

# ---------------------------------------------------------------------------
# PIPELINE DE CONVERSÃO PARA PDF
# ---------------------------------------------------------------------------
def run_lo_subprocess_with_backoff(cmd_args, env, max_retries=3, base_delay=0.5, max_delay=3.0):
    for attempt in range(max_retries):
        try:
            result = subprocess.run(
                cmd_args,
                capture_output=True,
                text=True,
                timeout=120,
                env=env,
            )
            if result.returncode == 0:
                return result
        except (subprocess.TimeoutExpired, OSError):
            pass

        if attempt < max_retries - 1:
            calculated_delay = min(max_delay, base_delay * (2 ** attempt))
            jitter = random.uniform(0, 0.3)
            time.sleep(calculated_delay + jitter)

    return None

def convert_to_pdf(input_file, output_dir):
    input_path = Path(input_file)
    profile_dir = Path(tempfile.gettempdir()) / f"lo_profile_{uuid.uuid4().hex}"
    ext = input_path.suffix.lower()

    try:
        # Pre-processamento inteligente de DOCX e XLSX
        if ext in [".docx", ".doc"]:
            sanitized_file = Path(tempfile.gettempdir()) / f"sanitized_{uuid.uuid4().hex}.docx"
            sanitize_docx_xml(str(input_path), str(sanitized_file))
            target_input = sanitized_file
        elif ext == ".xlsx":
            sanitized_file = Path(tempfile.gettempdir()) / f"sanitized_{uuid.uuid4().hex}.xlsx"
            sanitize_xlsx_page_setup(str(input_path), str(sanitized_file))
            target_input = sanitized_file
        else:
            sanitized_file = None
            target_input = input_path

        inject_lo_profile_settings(profile_dir)
        env = prepare_font_environment()

        result = None
        for cmd in ("soffice", "libreoffice"):
            cmd_args = [
                cmd,
                "--headless",
                "--norestore",
                f"-env:UserInstallation=file://{profile_dir}",
                "--convert-to", "pdf",
                "--outdir", str(output_dir),
                str(target_input),
            ]
            result = run_lo_subprocess_with_backoff(cmd_args, env=env)
            if result and result.returncode == 0:
                break

        if result is None or result.returncode != 0:
            error_msg = result.stderr.strip() if result and result.stderr else "Erro na conversão"
            st.error(f"❌ Erro na conversão: {error_msg}")
            return None

        pdf_generated = Path(output_dir) / (target_input.stem + ".pdf")
        final_pdf = Path(output_dir) / (input_path.stem + ".pdf")

        if pdf_generated.exists() and pdf_generated != final_pdf:
            pdf_generated.rename(final_pdf)

        if final_pdf.exists() and final_pdf.stat().st_size > 0:
            return str(final_pdf)

        return None

    except Exception as e:
        st.error(f"❌ Erro ao converter: {str(e)}")
        return None
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)
        if 'sanitized_file' in locals() and sanitized_file and sanitized_file.exists():
            try:
                sanitized_file.unlink()
            except OSError:
                pass

# ---------------------------------------------------------------------------
# MAIN INTERFACE
# ---------------------------------------------------------------------------
def main():
    uploaded_file = st.file_uploader(
        "Arraste e solte seu arquivo aqui",
        type=["doc", "docx", "xls", "xlsx", "ppt", "pptx", "odt", "ods", "odp"],
        help="Tamanho máximo: 200MB",
        label_visibility="collapsed",
    )

    if uploaded_file is None:
        return

    if (uploaded_file.size / (1024 * 1024)) > 200:
        st.error("❌ Arquivo muito grande! Máximo: 200MB")
        st.stop()

    with st.spinner("Convertendo documento para PDF..."):
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / uploaded_file.name
            input_path.write_bytes(uploaded_file.getbuffer())

            output_path = convert_to_pdf(str(input_path), temp_dir)
            if output_path is None:
                return

            output_bytes = Path(output_path).read_bytes()

    st.success("✅ Conversão concluída!")
    st.download_button(
        label="📥 Baixar PDF Convertido",
        data=output_bytes,
        file_name=Path(uploaded_file.name).stem + ".pdf",
        mime="application/pdf",
        type="primary",
        use_container_width=True,
    )

if __name__ == "__main__":
    main()
